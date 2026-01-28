#!/usr/bin/env python3
"""
Convert research-paper-for-AT.html to Word document format
"""
import re
from html.parser import HTMLParser
from html import unescape
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("python-docx not available. Install with: pip3 install python-docx --user")

class HTMLToWordParser(HTMLParser):
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_para = None
        self.in_research_paper = False
        self.in_para = False
        self.current_tag = None
        self.buffer = []
        self.list_level = 0
        
    def handle_starttag(self, tag, attrs):
        self.current_tag = tag
        attrs_dict = dict(attrs)
        
        # Check if entering research-paper div
        if tag == 'div' and ('class', 'research-paper') in attrs:
            self.in_research_paper = True
            return
            
        if not self.in_research_paper:
            return
            
        # Handle different tags
        if tag == 'h1':
            self.current_para = self.doc.add_heading('', level=1)
            self.current_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag == 'h2':
            self.current_para = self.doc.add_heading('', level=2)
        elif tag == 'h3':
            self.current_para = self.doc.add_heading('', level=3)
        elif tag == 'h4':
            self.current_para = self.doc.add_heading('', level=4)
        elif tag == 'p':
            self.current_para = self.doc.add_paragraph()
            # Check for center alignment
            if 'style' in attrs_dict and 'text-align: center' in attrs_dict['style']:
                self.current_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.in_para = True
        elif tag == 'strong':
            self.buffer.append(('strong', ''))
        elif tag == 'em':
            self.buffer.append(('em', ''))
        elif tag == 'sup':
            self.buffer.append(('sup', ''))
        elif tag == 'br':
            if self.current_para:
                self.current_para.add_run('\n')
        elif tag == 'li':
            if self.current_para:
                self.current_para.add_run('• ')
            else:
                self.current_para = self.doc.add_paragraph()
                self.current_para.add_run('• ')
        elif tag in ['ul', 'ol']:
            self.list_level += 1
        elif tag == 'table':
            self.current_table = None
        elif tag == 'tr':
            self.current_row = None
        elif tag in ['td', 'th']:
            self.current_cell = None
            
    def handle_endtag(self, tag):
        if tag == 'div' and self.in_research_paper:
            # Check if this is the closing research-paper div
            if self.current_tag == 'div':
                self.in_research_paper = False
            return
            
        if not self.in_research_paper:
            return
            
        if tag == 'p':
            self.in_para = False
            self.current_para = None
        elif tag == 'strong':
            self.buffer.append(('end_strong', ''))
        elif tag == 'em':
            self.buffer.append(('end_em', ''))
        elif tag == 'sup':
            self.buffer.append(('end_sup', ''))
        elif tag in ['h1', 'h2', 'h3', 'h4']:
            self.current_para = None
        elif tag in ['ul', 'ol']:
            self.list_level = max(0, self.list_level - 1)
            
    def handle_data(self, data):
        if not self.in_research_paper:
            return
            
        # Skip script and style content
        if self.current_tag in ['script', 'style', 'noscript']:
            return
            
        # Clean up data
        data = data.strip()
        if not data:
            return
            
        # Add text to current paragraph or create new one
        if not self.current_para:
            self.current_para = self.doc.add_paragraph()
            
        # Handle formatting
        run = self.current_para.add_run(data)
        
        # Apply formatting based on buffer
        for fmt, _ in self.buffer:
            if fmt == 'strong':
                run.bold = True
            elif fmt == 'em':
                run.italic = True
            elif fmt == 'sup':
                run.font.superscript = True
                
        # Clear buffer after processing
        self.buffer = []

def convert_html_to_word(html_file, output_file):
    """Convert HTML research paper to Word document"""
    
    if not HAS_DOCX:
        print("ERROR: python-docx is not installed.")
        print("Please install it with: pip3 install python-docx --user")
        return False
    
    # Read HTML file
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Extract content from research-paper div
    # Try to find the main content
    match = re.search(r'<div class="research-paper">(.*?)(?=</div>\s*</main>|$)', html_content, re.DOTALL)
    if not match:
        print("Could not find research-paper div")
        return False
        
    paper_html = match.group(1)
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Parse HTML
    parser = HTMLToWordParser(doc)
    parser.feed(paper_html)
    
    # Save document
    doc.save(output_file)
    print(f"Successfully created {output_file}")
    return True

if __name__ == '__main__':
    html_file = 'research-paper-for-AT.html'
    output_file = 'research-paper-for-AT.docx'
    
    if len(sys.argv) > 1:
        html_file = sys.argv[1]
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
        
    convert_html_to_word(html_file, output_file)
